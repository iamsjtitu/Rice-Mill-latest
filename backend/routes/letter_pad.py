"""
Letter Pad — Company letterhead generator with optional AI assistant.

Features:
  - Generate professional letterhead PDF + Word (.docx) from user-typed body
  - Letterhead matches the reference design: GSTIN top-left, ॐ + Company Name
    centered in red, address + email below; phone numbers top-right; red divider;
    Ref. No. left + Date right; body; signature block.
  - Optional AI assistant (Gemini 2.5 Flash OR GPT-5-mini) for:
      * Generate letter from a short prompt
      * Improve grammar/tone of typed text
      * Translate between English / Hindi / Odia
  - AI keys are per-installation (saved in app_settings) — each miller adds their
    own free Gemini key from Google AI Studio.
"""
from fastapi import APIRouter, HTTPException, Body
from fastapi.responses import StreamingResponse
from datetime import datetime, timezone
from io import BytesIO
import uuid
import logging
import httpx

from database import db
from utils.branding_helper import get_branding_data
from utils.export_helpers import register_hindi_fonts
from utils.letter_pad_templates import (
    LETTER_PAD_TEMPLATES,
    get_templates,
    get_template_by_id,
)
import re as _re

# Detect Devanagari (Hindi) characters → switch to Noto Devanagari font
_DEVA_RE = _re.compile(r"[\u0900-\u097F]")


def _has_deva(s) -> bool:
    return bool(_DEVA_RE.search(str(s or "")))


def _auto_font(text, bold: bool = False) -> str:
    """Pick Noto Devanagari for Hindi text, else Inter."""
    if _has_deva(text):
        return "NotoDevaBold" if bold else "NotoDeva"
    return "InterBold" if bold else "Inter"

router = APIRouter()
logger = logging.getLogger("letter_pad")

# ---------- Settings (signature + AI keys) ----------

@router.get("/letter-pad/settings")
async def get_letter_pad_settings():
    """Return signature + letterhead fields + AI key presence (never the actual keys)."""
    doc = await db.app_settings.find_one({"setting_id": "letter_pad"}, {"_id": 0}) or {}
    return {
        "gstin": doc.get("gstin", ""),
        "phone": doc.get("phone", ""),
        "phone_secondary": doc.get("phone_secondary", ""),
        "address": doc.get("address", ""),
        "email": doc.get("email", ""),
        "license_number": doc.get("license_number", ""),
        "header_text": doc.get("header_text", ""),
        "signature_name": doc.get("signature_name", ""),
        "signature_designation": doc.get("signature_designation", ""),
        "ai_enabled": bool(doc.get("ai_enabled", False)),
        "has_gemini_key": bool(doc.get("gemini_key")),
        "has_openai_key": bool(doc.get("openai_key")),
        "ai_provider": doc.get("ai_provider", "gemini"),
    }


@router.put("/letter-pad/settings")
async def update_letter_pad_settings(payload: dict = Body(...)):
    """Update letterhead + signature + AI configuration. Empty key strings are
    ignored (so users can update other fields without re-entering keys).
    """
    update = {}
    text_fields = (
        "gstin", "phone", "phone_secondary", "address", "email", "license_number",
        "signature_name", "signature_designation", "ai_provider", "header_text",
    )
    for f in text_fields:
        if f in payload:
            update[f] = str(payload.get(f) or "").strip()
    if "ai_enabled" in payload:
        update["ai_enabled"] = bool(payload["ai_enabled"])
    if payload.get("gemini_key"):
        update["gemini_key"] = str(payload["gemini_key"]).strip()
    if payload.get("openai_key"):
        update["openai_key"] = str(payload["openai_key"]).strip()
    if payload.get("clear_gemini_key"):
        update["gemini_key"] = ""
    if payload.get("clear_openai_key"):
        update["openai_key"] = ""
    update["updated_at"] = datetime.utcnow().isoformat()
    await db.app_settings.update_one(
        {"setting_id": "letter_pad"},
        {"$set": update, "$setOnInsert": {"setting_id": "letter_pad"}},
        upsert=True,
    )
    return await get_letter_pad_settings()


# ---------- AI proxy ----------

_AI_SYSTEM_PROMPTS = {
    "generate": (
        "You are a professional business letter writer. Generate a complete formal "
        "Indian business letter and return STRICT JSON output with three keys.\n\n"
        "OUTPUT FORMAT (return ONLY this JSON, no markdown fences, no preamble):\n"
        "{\n"
        '  "to_address": "<recipient name and address, multi-line with \\n>",\n'
        '  "subject": "<concise subject line, 5-12 words>",\n'
        '  "body": "<full letter body starting with greeting, ending with Thanking you.>"\n'
        "}\n\n"
        "STRICT RULES for the body:\n"
        "1. NO sender's company info (it's already on the letterhead).\n"
        "2. NO placeholders like '[Your Name]', '[Date]', '[Recipient]'.\n"
        "3. Start with 'Respected Sir/Madam,' (or appropriate greeting).\n"
        "4. End with 'Thanking you.' (NO 'Yours faithfully' / signature).\n"
        "5. Use first person plural (we/our).\n"
        "6. 150-300 words. Polite, direct, formal.\n\n"
        "STRICT RULES for to_address:\n"
        "- Standard Indian business format. Example:\n"
        "  'The Branch Manager,\\nState Bank of India,\\n[Branch Name],\\n[City]'\n"
        "- If user did not specify recipient, infer reasonable defaults based on context.\n\n"
        "Match the language requested (Hindi/English/Odia) for ALL three fields.\n"
        "Return PURE JSON only. No ```json fences."
    ),
    "improve": (
        "You are a professional business letter editor. Rewrite the user's draft "
        "with improved grammar, tone, and professionalism. "
        "STRICT RULES:\n"
        "1. Output ONLY the improved letter body — no preamble, no explanation.\n"
        "2. NO placeholders like '[Your Name]'.\n"
        "3. Preserve the user's intent and key facts (dates, amounts, names).\n"
        "4. Don't add 'Yours faithfully' / signature — those are added separately.\n"
        "5. Keep length similar (don't pad or shorten drastically). Same language as input."
    ),
    "translate": (
        "You are a professional translator for Indian business letters. Translate "
        "between English, Hindi (Devanagari), and Odia. "
        "STRICT RULES:\n"
        "1. Output ONLY the translation — no preamble, no explanation.\n"
        "2. Preserve formal business tone and all factual details (numbers, dates, names).\n"
        "3. Use natural, professional phrasing in the target language."
    ),
}


async def _call_gemini(api_key: str, system_prompt: str, user_prompt: str, json_mode: bool = False) -> str:
    url = (
        "https://generativelanguage.googleapis.com/v1beta/models/"
        f"gemini-2.5-flash:generateContent?key={api_key}"
    )
    gen_config = {
        "temperature": 0.7,
        "maxOutputTokens": 2048,
        "thinkingConfig": {"thinkingBudget": 0},
    }
    if json_mode:
        gen_config["responseMimeType"] = "application/json"
    payload = {
        "system_instruction": {"parts": [{"text": system_prompt}]},
        "contents": [{"role": "user", "parts": [{"text": user_prompt}]}],
        "generationConfig": gen_config,
    }
    async with httpx.AsyncClient(timeout=60) as client:
        r = await client.post(url, json=payload)
        if r.status_code != 200:
            raise HTTPException(status_code=502, detail=f"Gemini error: {r.text[:300]}")
        data = r.json()
        try:
            return data["candidates"][0]["content"]["parts"][0]["text"].strip()
        except (KeyError, IndexError):
            raise HTTPException(status_code=502, detail="Gemini returned no text")


async def _call_openai(api_key: str, system_prompt: str, user_prompt: str, json_mode: bool = False) -> str:
    url = "https://api.openai.com/v1/chat/completions"
    payload = {
        "model": "gpt-5-mini",
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "max_completion_tokens": 1500,
    }
    if json_mode:
        payload["response_format"] = {"type": "json_object"}
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    async with httpx.AsyncClient(timeout=60) as client:
        r = await client.post(url, json=payload, headers=headers)
        if r.status_code != 200:
            raise HTTPException(status_code=502, detail=f"OpenAI error: {r.text[:300]}")
        data = r.json()
        try:
            return data["choices"][0]["message"]["content"].strip()
        except (KeyError, IndexError):
            raise HTTPException(status_code=502, detail="OpenAI returned no text")


@router.post("/letter-pad/ai")
async def letter_pad_ai(payload: dict = Body(...)):
    """payload = { mode: 'generate'|'improve'|'translate', text: '...',
                   target_lang?: 'English'|'Hindi'|'Odia' }"""
    mode = payload.get("mode")
    text = (payload.get("text") or "").strip()
    if mode not in _AI_SYSTEM_PROMPTS:
        raise HTTPException(status_code=400, detail="Invalid mode")
    if not text:
        raise HTTPException(status_code=400, detail="Text/prompt is required")
    settings = await db.app_settings.find_one({"setting_id": "letter_pad"}, {"_id": 0}) or {}
    if not settings.get("ai_enabled"):
        raise HTTPException(status_code=400, detail="AI is disabled. Settings → Letter Pad → AI Assistant on karein.")
    provider = settings.get("ai_provider", "gemini")
    gemini_key = settings.get("gemini_key", "")
    openai_key = settings.get("openai_key", "")
    if provider == "gemini" and not gemini_key:
        if openai_key:
            provider = "openai"
        else:
            raise HTTPException(status_code=400, detail="Gemini API key not set")
    elif provider == "openai" and not openai_key:
        if gemini_key:
            provider = "gemini"
        else:
            raise HTTPException(status_code=400, detail="OpenAI API key not set")

    system_prompt = _AI_SYSTEM_PROMPTS[mode]
    user_prompt = text
    if mode == "translate":
        target = payload.get("target_lang") or "English"
        user_prompt = f"Translate the following to {target}:\n\n{text}"

    json_mode = (mode == "generate")
    if provider == "gemini":
        out = await _call_gemini(gemini_key, system_prompt, user_prompt, json_mode=json_mode)
    else:
        out = await _call_openai(openai_key, system_prompt, user_prompt, json_mode=json_mode)

    # For generate mode, parse JSON and return structured fields
    if json_mode:
        import json as _json
        import re as _re
        clean = out.strip()
        # Strip markdown fences if model added them despite instruction
        clean = _re.sub(r"^```(?:json)?\s*", "", clean)
        clean = _re.sub(r"\s*```$", "", clean)
        try:
            parsed = _json.loads(clean)
            return {
                "result": parsed.get("body", "").strip(),
                "subject": parsed.get("subject", "").strip(),
                "to_address": parsed.get("to_address", "").strip(),
                "provider": provider,
                "structured": True,
            }
        except Exception:
            # Fallback: treat whole output as body
            return {"result": clean, "subject": "", "to_address": "", "provider": provider, "structured": False}

    return {"result": out, "provider": provider}


# ---------- PDF letterhead ----------

async def _build_letter_context():
    """Pull branding + letter pad customization from DB. Letter Pad's own
    overrides (GSTIN, phone, email, address, license) take priority over the
    global branding values."""
    branding = await get_branding_data()
    lp = await db.app_settings.find_one({"setting_id": "letter_pad"}, {"_id": 0}) or {}

    def pick(*opts):
        for o in opts:
            if o:
                return o
        return ""

    return {
        "company_name": pick(branding.get("company_name"), "NAVKAR AGRO"),
        "tagline": pick(branding.get("tagline")),
        "header_text": pick(lp.get("header_text")),
        "address": pick(lp.get("address"), branding.get("address"), "Laitara Road, Jolko - 766012, Dist. Kalahandi (Odisha)"),
        "email": pick(lp.get("email"), branding.get("email")),
        "phone": pick(lp.get("phone"), branding.get("phone")),
        "phone_secondary": pick(lp.get("phone_secondary"), branding.get("phone_secondary")),
        "gstin": pick(lp.get("gstin"), branding.get("gstin")),
        "license_number": pick(lp.get("license_number"), branding.get("mill_code")),
        "logo": pick(branding.get("logo")),
        "signature_name": pick(lp.get("signature_name"), "Aditya Jain"),
        "signature_designation": pick(lp.get("signature_designation"), "Proprietor"),
    }


def _draw_letterhead_pdf(canvas, ctx, page_w, page_h):
    """Draw the Navkar-style letterhead on the given canvas (top portion).

    Layout:
      Row 1: GSTIN top-left  |  Mob phones top-right
      Row 2: Centered ॐ + Company Name (red bold, 22pt)
      Row 3: Address (centered, 10pt slate)
      Row 4: Email (centered, 10pt slate)
      Row 5: Red divider line
      Row 6 (optional): License number centered below divider, 9pt
    Returns the y position where body content can start.
    """
    from reportlab.lib.colors import HexColor

    register_hindi_fonts()
    BRAND_RED = HexColor("#C0392B")
    DARK = HexColor("#1f2937")
    MUTED = HexColor("#475569")

    top = page_h - 38

    # GSTIN top-left
    if ctx["gstin"]:
        canvas.setFont("InterBold", 9)
        canvas.setFillColor(DARK)
        canvas.drawString(40, top, f"GSTIN: {ctx['gstin']}")

    # Phone(s) top-right (one per line)
    canvas.setFont("InterBold", 9)
    canvas.setFillColor(DARK)
    phones = []
    if ctx["phone"]:
        phones.append(ctx["phone"])
    if ctx["phone_secondary"]:
        phones.append(ctx["phone_secondary"])
    for i, p in enumerate(phones):
        canvas.drawRightString(page_w - 40, top - i * 11, f"Mob. {p}")

    # Center: optional small header text (slogan like "Shree Ram") ABOVE company name
    company_y = top - 28
    if ctx.get("header_text"):
        canvas.setFont(_auto_font(ctx["header_text"], bold=False), 11)
        canvas.setFillColor(MUTED)
        canvas.drawCentredString(page_w / 2, top - 12, ctx["header_text"])
        company_y = top - 36

    # Center: Company Name (red, bold). Pulled down a bit so phone numbers fit.
    canvas.setFont(_auto_font(ctx['company_name'], bold=True), 22)
    canvas.setFillColor(BRAND_RED)
    canvas.drawCentredString(page_w / 2, company_y, ctx['company_name'])

    # Address (centered)
    addr_y = top - 46
    canvas.setFont("Inter", 10)
    canvas.setFillColor(MUTED)
    if ctx["address"]:
        canvas.drawCentredString(page_w / 2, addr_y, ctx["address"])
        addr_y -= 13

    # Email (centered)
    if ctx["email"]:
        canvas.setFont("Inter", 10)
        canvas.setFillColor(MUTED)
        canvas.drawCentredString(page_w / 2, addr_y, f"Email: {ctx['email']}")
        addr_y -= 13

    # Red divider line
    divider_y = addr_y - 6
    canvas.setStrokeColor(BRAND_RED)
    canvas.setLineWidth(1.5)
    canvas.line(40, divider_y, page_w - 40, divider_y)

    # License number (optional, below divider, centered, small)
    if ctx.get("license_number"):
        canvas.setFont("Inter", 8)
        canvas.setFillColor(MUTED)
        canvas.drawCentredString(page_w / 2, divider_y - 11, f"License No: {ctx['license_number']}")
        return divider_y - 22
    return divider_y - 8


def _draw_footer_signature(canvas, ctx, page_w, y):
    """Draw 'Yours faithfully,' + signature name + designation. Tighter spacing."""
    canvas.setFont("Inter", 11)
    canvas.setFillColor("#1f2937")
    canvas.drawRightString(page_w - 40, y, "Yours faithfully,")
    canvas.setFont(_auto_font(ctx["signature_name"], bold=True), 12)
    canvas.drawRightString(page_w - 40, y - 32, ctx["signature_name"])
    canvas.setFont("Inter", 10)
    canvas.setFillColor("#475569")
    canvas.drawRightString(page_w - 40, y - 46, ctx["signature_designation"])
    canvas.setFont(_auto_font(ctx["company_name"], bold=False), 10)
    canvas.drawRightString(page_w - 40, y - 60, f"M/s {ctx['company_name']}")


@router.post("/letter-pad/pdf")
async def generate_letter_pdf(payload: dict = Body(...)):
    """payload = { ref_no, date, to_address, subject, body, references }"""
    ctx = await _build_letter_context()
    pdf_bytes = _build_letter_pdf_bytes(payload, ctx)
    fname = f"letter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    return StreamingResponse(
        BytesIO(pdf_bytes), media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


def _build_letter_pdf_bytes(payload: dict, ctx: dict) -> bytes:
    """Render letter to PDF bytes (used by both /pdf endpoint and WhatsApp share)."""
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.pagesizes import A4

    buf = BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=A4)
    page_w, page_h = A4

    # === Top letterhead — returns y where content starts ===
    content_top_y = _draw_letterhead_pdf(c, ctx, page_w, page_h)

    y = content_top_y - 18
    c.setFillColor("#1f2937")
    c.setFont("Inter", 10)
    # Ref. No. (left) + Date (right)
    ref_no = (payload.get("ref_no") or "").strip()
    date = (payload.get("date") or datetime.now().strftime("%d-%m-%Y")).strip()
    c.drawString(40, y, f"Ref. No.: {ref_no or '_____________'}")
    c.drawRightString(page_w - 40, y, f"Date: {date}")
    y -= 30

    # === To ===
    to_addr = (payload.get("to_address") or "").strip()
    if to_addr:
        c.setFont("InterBold", 10)
        c.drawString(40, y, "To,")
        y -= 14
        for line in to_addr.split("\n"):
            c.setFont(_auto_font(line, bold=False), 10)
            c.drawString(50, y, line[:90])
            y -= 13
        y -= 8

    # === Subject ===
    subject = (payload.get("subject") or "").strip()
    if subject:
        c.setFont(_auto_font(subject, bold=True), 11)
        c.drawString(40, y, f"Subject: {subject}")
        y -= 18

    # === References (optional) ===
    refs = (payload.get("references") or "").strip()
    if refs:
        c.setFont("InterBold", 10)
        c.drawString(40, y, "Reference:")
        y -= 13
        for line in refs.split("\n"):
            c.setFont(_auto_font(line, bold=False), 10)
            c.drawString(50, y, line[:95])
            y -= 13
        y -= 6

    # === Body (greeting + paragraphs) ===
    from reportlab.lib.colors import HexColor as _HexColor
    body = (payload.get("body") or "").strip()
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Paragraph
    from reportlab.lib.enums import TA_JUSTIFY
    body_font = "NotoDeva" if _has_deva(body) else "Inter"
    c.setFont(body_font, 11)
    body_style = ParagraphStyle(
        "Body", fontName=body_font, fontSize=11, leading=15,
        alignment=TA_JUSTIFY, textColor=_HexColor("#1f2937"),
    )
    for para in body.split("\n\n"):
        if not para.strip():
            continue
        p = Paragraph(para.replace("\n", "<br/>"), body_style)
        avail_h = y - 150
        w, h = p.wrap(page_w - 80, avail_h)
        if h > avail_h:
            c.showPage()
            new_top = _draw_letterhead_pdf(c, ctx, page_w, page_h)
            y = new_top - 18
            w, h = p.wrap(page_w - 80, page_h - 200)
        p.drawOn(c, 40, y - h)
        y -= h + 8

    # === Signature ===
    # Signature 30pt neeche jaha body khatam hua. Agar body bahut lambi hai
    # (signature page ke neeche overflow karega), to bottom margin pe clamp.
    sig_y = max(120, y - 30)
    _draw_footer_signature(c, ctx, page_w, sig_y)

    c.save()
    return buf.getvalue()



# ---------- Word (.docx) letterhead ----------

@router.post("/letter-pad/docx")
async def generate_letter_docx(payload: dict = Body(...)):
    """Generate a Word document with the same letterhead. Editable in MS Word."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ctx = await _build_letter_context()
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # === Letterhead header (table for layout) ===
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    cells = table.rows[0].cells
    if ctx["gstin"]:
        p = cells[0].paragraphs[0]
        run = p.add_run(f"GSTIN: {ctx['gstin']}")
        run.font.size = Pt(9)
    p_right = cells[1].paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if ctx["phone"]:
        run = p_right.add_run(f"Mob. {ctx['phone']}")
        run.font.size = Pt(9)
    if ctx["phone_secondary"]:
        p2 = cells[1].add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p2.add_run(f"Mob. {ctx['phone_secondary']}")
        run.font.size = Pt(9)

    # Center: optional small header text (slogan) ABOVE company name
    if ctx.get("header_text"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(ctx["header_text"])
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    # Center: company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(ctx['company_name'])
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    if ctx["address"]:
        p = doc.add_paragraph(ctx["address"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    if ctx["email"]:
        p = doc.add_paragraph(f"Email: {ctx['email']}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    # Red divider (horizontal rule via bottom border on a paragraph)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C0392B")
    border.append(bottom)
    p_pr.append(border)

    # === Ref + Date row ===
    table = doc.add_table(rows=1, cols=2)
    cells = table.rows[0].cells
    p = cells[0].paragraphs[0]
    p.add_run(f"Ref. No.: {payload.get('ref_no') or '_____________'}").font.size = Pt(10)
    p_r = cells[1].paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_str = payload.get("date") or datetime.now().strftime("%d-%m-%Y")
    p_r.add_run(f"Date: {date_str}").font.size = Pt(10)

    doc.add_paragraph()

    # === To ===
    if payload.get("to_address"):
        p = doc.add_paragraph()
        run = p.add_run("To,")
        run.bold = True
        run.font.size = Pt(11)
        for line in payload["to_address"].split("\n"):
            sub = doc.add_paragraph(line)
            sub.paragraph_format.left_indent = Cm(0.5)
            for r in sub.runs:
                r.font.size = Pt(10)

    # === Subject ===
    if payload.get("subject"):
        p = doc.add_paragraph()
        run = p.add_run(f"Subject: {payload['subject']}")
        run.bold = True
        run.font.size = Pt(11)

    # === References (optional) ===
    if payload.get("references"):
        p = doc.add_paragraph()
        run = p.add_run("Reference:")
        run.bold = True
        run.font.size = Pt(10)
        for line in payload["references"].split("\n"):
            sub = doc.add_paragraph(line)
            sub.paragraph_format.left_indent = Cm(0.5)
            for r in sub.runs:
                r.font.size = Pt(10)

    # === Body ===
    body = (payload.get("body") or "").strip()
    for para in body.split("\n\n"):
        if not para.strip():
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        for line in para.split("\n"):
            run = p.add_run(line)
            run.font.size = Pt(11)
            run.add_break()

    # === Signature ===
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Yours faithfully,").font.size = Pt(11)
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(ctx["signature_name"])
    run.bold = True
    run.font.size = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(ctx["signature_designation"])
    run.font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"M/s {ctx['company_name']}")
    run.font.size = Pt(10)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    fname = f"letter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )



# ====================== DRAFTS (CRUD) ======================

def _draft_doc(d: dict) -> dict:
    """Strip MongoDB internals."""
    d.pop("_id", None)
    return d


@router.get("/letter-pad/drafts")
async def list_drafts():
    """List saved drafts (newest first)."""
    cursor = db.letter_drafts.find({}, {"_id": 0}).sort("updated_at", -1)
    drafts = []
    async for d in cursor:
        drafts.append(d)
    return drafts


@router.post("/letter-pad/drafts")
async def create_draft(payload: dict = Body(...)):
    """Create a new draft. payload = {title?, ref_no, date, to_address, subject,
    references, body}."""
    body = (payload.get("body") or "").strip()
    if not body and not (payload.get("subject") or "").strip():
        raise HTTPException(status_code=400, detail="Khaali draft save nahi ho sakti — kuch text type karein")
    now = datetime.now(timezone.utc).isoformat()
    title = (payload.get("title") or "").strip() or (payload.get("subject") or "").strip()[:60] or "Untitled Draft"
    doc = {
        "id": str(uuid.uuid4()),
        "title": title,
        "ref_no": payload.get("ref_no", ""),
        "date": payload.get("date", ""),
        "to_address": payload.get("to_address", ""),
        "subject": payload.get("subject", ""),
        "references": payload.get("references", ""),
        "body": payload.get("body", ""),
        "created_at": now,
        "updated_at": now,
    }
    await db.letter_drafts.insert_one(doc.copy())
    return _draft_doc(doc)


@router.put("/letter-pad/drafts/{draft_id}")
async def update_draft(draft_id: str, payload: dict = Body(...)):
    existing = await db.letter_drafts.find_one({"id": draft_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Draft not found")
    update_fields = {}
    for f in ("title", "ref_no", "date", "to_address", "subject", "references", "body"):
        if f in payload:
            update_fields[f] = payload[f] or ""
    update_fields["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db.letter_drafts.update_one({"id": draft_id}, {"$set": update_fields})
    return await db.letter_drafts.find_one({"id": draft_id}, {"_id": 0})


@router.delete("/letter-pad/drafts/{draft_id}")
async def delete_draft(draft_id: str):
    res = await db.letter_drafts.delete_one({"id": draft_id})
    if res.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Draft not found")
    return {"success": True}


# ====================== TEMPLATES ======================

@router.get("/letter-pad/templates")
async def list_templates():
    """Lightweight list (id, name, category) for picker UI."""
    return get_templates()


@router.get("/letter-pad/templates/{template_id}")
async def get_template(template_id: str):
    t = get_template_by_id(template_id)
    if not t:
        raise HTTPException(status_code=404, detail="Template not found")
    return t


# ====================== WHATSAPP SHARE ======================

@router.post("/letter-pad/whatsapp")
async def share_letter_via_whatsapp(payload: dict = Body(...)):
    """Generate letter PDF, upload to tmpfiles.org, send via 360Messenger.

    payload = {
      letter: {ref_no, date, to_address, subject, references, body},
      mode: 'phone' | 'group' | 'default',
      phone?: '9876543210',
      group_id?: '...',
      caption?: 'Custom message'
    }
    """
    from routes.whatsapp import _get_wa_settings, _send_wa_message, _send_wa_to_group

    letter = payload.get("letter") or {}
    if not (letter.get("body") or "").strip():
        raise HTTPException(status_code=400, detail="Letter body khaali hai")
    mode = (payload.get("mode") or "default").lower()

    wa_settings = await _get_wa_settings()
    if not wa_settings.get("api_key"):
        raise HTTPException(status_code=400, detail="WhatsApp API key set nahi hai. Settings → WhatsApp mein 360Messenger key dale.")

    # Build PDF in memory
    ctx = await _build_letter_context()
    pdf_bytes = _build_letter_pdf_bytes(letter, ctx)
    if len(pdf_bytes) < 100:
        raise HTTPException(status_code=500, detail="PDF generation fail")

    # Upload to tmpfiles.org
    public_pdf_url = ""
    try:
        async with httpx.AsyncClient(timeout=30) as client:
            files = {"file": (f"letter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf", pdf_bytes, "application/pdf")}
            up = await client.post("https://tmpfiles.org/api/v1/upload", files=files)
            if up.status_code == 200:
                tmp_url = up.json().get("data", {}).get("url", "")
                if tmp_url:
                    public_pdf_url = tmp_url.replace("http://tmpfiles.org/", "https://tmpfiles.org/dl/")
    except Exception as e:
        logger.error(f"Letter PDF upload error: {e}")
    if not public_pdf_url:
        raise HTTPException(status_code=502, detail="PDF upload (tmpfiles.org) fail hua. Internet check karein.")

    # Caption: company name + subject + custom note
    company = ctx.get("company_name") or "Mill Entry System"
    subject = (letter.get("subject") or "").strip()
    note = (payload.get("caption") or "").strip()
    caption = f"*{company}*"
    if subject:
        caption += f"\nSubject: {subject}"
    if note:
        caption += f"\n{note}"
    else:
        caption += "\nPlease find attached letter."
    caption += f"\n\n— {company}"

    results = []
    if mode == "phone":
        phone = (payload.get("phone") or "").strip()
        if not phone:
            raise HTTPException(status_code=400, detail="Phone number daalein")
        r = await _send_wa_message(phone, caption, public_pdf_url)
        results.append({"target": phone, "success": r.get("success", False), "error": r.get("error", "")})
    elif mode == "group":
        gid = (payload.get("group_id") or "").strip() or wa_settings.get("default_group_id", "") or wa_settings.get("group_id", "")
        if not gid:
            raise HTTPException(status_code=400, detail="Group ID daalein ya default group set karein")
        r = await _send_wa_to_group(gid, caption, public_pdf_url)
        results.append({"target": "group", "success": r.get("success", False), "error": r.get("error", "")})
    else:  # default — send to all default numbers
        nums = wa_settings.get("default_numbers", [])
        if isinstance(nums, str):
            nums = [n.strip() for n in nums.split(",") if n.strip()]
        if not nums:
            raise HTTPException(status_code=400, detail="Default numbers set nahi hai. Settings → WhatsApp mein numbers SAVE karein, ya phone/group choose karein.")
        for num in nums:
            r = await _send_wa_message(num, caption, public_pdf_url)
            results.append({"target": num, "success": r.get("success", False), "error": r.get("error", "")})

    success_count = sum(1 for r in results if r["success"])
    return {
        "success": success_count > 0,
        "message": f"Letter {success_count}/{len(results)} target(s) pe bhej diya!",
        "details": results,
        "pdf_url": public_pdf_url,
    }
